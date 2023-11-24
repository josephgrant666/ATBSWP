import docx
# Word runs three types of data objects. 1) For the doc itself 2) Paragraph objects and 3) Run objects 
d = docx.Document('Filename') # This returns a doc object
d.paragraphs # Returns information on all of the paragraphs objects in the doc 
d.paragraphs[0].text # Returns a string of the text within the paragraph according to the index 
p = d.paragraphs[0] # Saves the paragraph object 
p.runs # Returns a list of run objects. A run is whenever there is a change in style of the text; e.g. bold/italics
p.runs[0].text # Returns a string according to the index of the run 
p.runs[1].bold # Returns whether or not the run is bold or not  
p.runs[1].italic # Returns whether or not the run is italic or not  
p.runs[1].underline # Returns whether or not the run is underlined or not  
d.save('filepath') # Saves the file to the chosen path
p.style # Returns the style of the selected paragraph
p.style = 'Title' # Changes the style to a title
d.add_paragraph('Hello this is a paragraph.') # This adds a paragraph to the document object 
p.add_run('This is a new run.') # This adds a new run to the doc object 
